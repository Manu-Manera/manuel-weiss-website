"""
Word-Vorlagen (.docx): Platzhalter ersetzen und nach HTML konvertieren (mammoth).
"""

from __future__ import annotations

import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.table import Table


def _apply_mapping_to_paragraph(paragraph, mapping: dict[str, str]) -> None:
    text = paragraph.text
    newt = text
    for key, val in mapping.items():
        newt = newt.replace(key, val)
    if newt == text:
        return
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = newt
    else:
        paragraph.add_run(newt)


def _walk_table(table: Table, mapping: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _apply_mapping_to_paragraph(p, mapping)
            for nested in cell.tables:
                _walk_table(nested, mapping)


def _walk_header_footer(part, mapping: dict[str, str]) -> None:
    if part is None:
        return
    for p in part.paragraphs:
        _apply_mapping_to_paragraph(p, mapping)
    for table in part.tables:
        _walk_table(table, mapping)


def replace_placeholders_in_docx(path: Path, mapping: dict[str, str]) -> None:
    """
    Ersetzt Platzhalter-Strings (z. B. {NAME}) in allen Absätzen inkl. Tabellen und Kopf-/Fußzeilen.
    """
    doc = Document(str(path))
    for p in doc.paragraphs:
        _apply_mapping_to_paragraph(p, mapping)
    for table in doc.tables:
        _walk_table(table, mapping)
    for section in doc.sections:
        _walk_header_footer(section.header, mapping)
        if section.different_first_page_header_footer:
            _walk_header_footer(section.first_page_header, mapping)
        _walk_header_footer(section.footer, mapping)
        if section.different_first_page_header_footer:
            _walk_header_footer(section.first_page_footer, mapping)
    doc.save(str(path))


def docx_to_filled_html(template_path: Path, mapping: dict[str, str]) -> str:
    """Kopiert die Vorlage, ersetzt Platzhalter, konvertiert mit mammoth nach HTML."""
    import mammoth

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        shutil.copy2(template_path, tmp_path)
        replace_placeholders_in_docx(tmp_path, mapping)
        with open(tmp_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        return result.value
    finally:
        tmp_path.unlink(missing_ok=True)


def export_docx_filled_copy(
    template_path: Path,
    output_path: Path,
    mapping: dict[str, str],
) -> None:
    """Schreibt eine gefüllte Kopie der Word-Vorlage nach output_path."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    replace_placeholders_in_docx(output_path, mapping)
