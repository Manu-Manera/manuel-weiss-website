"""Export E-Mail-Vorlagentext als Word-Datei (.docx)."""

from __future__ import annotations

import re
from pathlib import Path


def _resolve_img_src_for_docx(html: str, template_root: Path | None) -> str:
    """Replace relative img src with file:// URIs so html4docx can embed images."""
    if not template_root or not template_root.is_dir():
        return html
    root = template_root.resolve()

    def repl(m: re.Match) -> str:
        quote = m.group(1)
        src = m.group(2).strip()
        if src.startswith(("http://", "https://", "data:", "file:")):
            return m.group(0)
        try:
            p = (root / src).resolve()
        except (OSError, ValueError):
            return m.group(0)
        if not p.is_file():
            return m.group(0)
        if root != p and root not in p.parents:
            return m.group(0)
        return f"src={quote}{p.as_uri()}{quote}"

    return re.sub(r"src\s*=\s*(['\"])([^'\"]+)\1", repl, html, flags=re.IGNORECASE)


def export_body_to_docx(
    *,
    body: str,
    is_html: bool,
    output_path: Path,
    title: str = "",
    template_root: Path | None = None,
) -> None:
    """
    Schreibt den Vorlagen-Inhalt nach output_path (.docx).

    - HTML: Konvertierung über html4docx (Absätze, Fett, Listen, … je nach Library).
    - Text: ein Absatz pro Zeile.
    """
    from docx import Document

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    head = (title or "").strip()
    if head:
        doc.add_heading(head, level=1)

    if is_html and body.strip():
        from html4docx import HtmlToDocx

        html = _resolve_img_src_for_docx(body, template_root)
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)
    else:
        for line in body.splitlines():
            doc.add_paragraph(line if line.strip() else " ")

    doc.save(str(output_path))
