#!/usr/bin/env python3
"""Knauf PM/SM Training aus der Almirall-Word-Vorlage — **tiefe** Transformation.

Siehe ``PROMPT_UND_ANFORDERUNGEN_KNAUF_PM.md`` für die vollständige Spezifikation.

Kernpunkte dieser Version:
    • Struktur-Anpassungen wie zuvor (BPA entfernen, Financials einfügen, Kapitelreihenfolge).
    • **OPC-ZIP-Gesamt-Ersatz** in allen ``word/**/*.xml`` — erwischt Hyperlinks/Fußzeilen/
      Split-Runs, die python-docx nicht erreicht.
    • **Automatischer Bild-Tausch** aller nicht‑Mini-Grafiken gegen einen Pool aus den
      bestehenden Knauf-Guides (aspect-fit wie SHS Adapted).
    • Optionales Überschreiben einzelner Slots über ``config/image_mapping.json``.
"""

from __future__ import annotations

import io
import json
import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as DocxParagraph
from PIL import Image


HERE = Path(__file__).resolve().parent
DOCS_ROOT = HERE.parent  # Onboarding Valkeen/docs

TEMPLATE = HERE / "Almirall_PM_Training_ORIGINAL.docx"
OUTPUT_DOCX = HERE / "Knauf_User_Training_PM_Almirall_Format.docx"
CONFIG_DIR = HERE / "config"
SCREENSHOTS_LOCAL = HERE / "screenshots"

KNAUF_TEMPUS_URL = "https://knaufdev.tempus-resource.com/"


# Reihenfolge: längere / spezifischere Strings zuerst (Substring-Ersatz).
OPC_XML_TEXT_REPLACEMENTS: list[tuple[str, str]] = [
    ("Almirall Hermal", "Knauf SE"),
    ("https://almirall.tempus-resource.com/sg/", KNAUF_TEMPUS_URL),
    ("http://almirall.tempus-resource.com/sg/", KNAUF_TEMPUS_URL),
    ("almirall.tempus-resource.com", "knaufdev.tempus-resource.com"),
    ("Almirall", "Knauf"),
    ("almirall", "knauf"),
    ("ALMIRALL", "KNAUF"),
    ("Continue With Microsoft", "Knauf SSO"),
]


def discover_knauf_screenshot_pool() -> list[Path]:
    """Sammelt PNG aus allen Knauf-Guide-Verzeichnissen unter ``docs/``."""
    candidates: list[Path] = []
    globs = (
        "[0-9][0-9]_*.png",
        "qrg_*.png",
        "details/*.png",
        "annotated/*_annot.png",
    )
    roots = [
        DOCS_ROOT / "Knauf PM SM Guide/screenshots",
        DOCS_ROOT / "Knauf QRG v2/screenshots",
        DOCS_ROOT / "Knauf User Training/screenshots",
        DOCS_ROOT / "Knauf Quick Reference Guide/screenshots",
        DOCS_ROOT / "Knauf SHS Adapted/screenshots",
    ]
    for root in roots:
        if not root.exists():
            continue
        for g in globs:
            candidates.extend(sorted(root.glob(g)))
    # Dedupe bei stabiler Reihenfolge
    seen: set[str] = set()
    uniq: list[Path] = []
    for p in candidates:
        if not p.is_file():
            continue
        key = str(p.resolve())
        if key not in seen:
            seen.add(key)
            uniq.append(p)
    if not uniq:
        raise RuntimeError(
            "Kein Knauf-Screenshot-Pool gefunden — erwartet PNGs unter Knauf PM SM Guide/, "
            "Knauf QRG v2/, Knauf User Training/, Knauf Quick Reference Guide/, Knauf SHS Adapted/."
        )
    return uniq


def _media_dimensions(data: bytes) -> tuple[int, int]:
    try:
        im = Image.open(io.BytesIO(data))
        return im.size
    except Exception:
        return (0, 0)


def should_replace_embedded_media(filename: str, data: bytes) -> bool:
    """Mini-Icons behalten (tenant-neutral); größere Einbettungen ersetzen."""
    lower = filename.lower()
    if not (
        lower.endswith(".png") or lower.endswith(".jpg") or lower.endswith(".jpeg")
    ):
        return False
    w, h = _media_dimensions(data)
    if w <= 0:
        return False
    area = w * h
    longest = max(w, h)
    shortest = min(w, h)
    kb = len(data) // 1024
    # Kleine Toolbar-/Glyph-Grafiken typischerweise <=72 px und wenig Pixelfläche
    if longest <= 72 and area <= 72 * 72:
        return False
    # Breite aber sehr flache Chrome-Leisten (Pixelhöhe klein, aber sehr breit): trotzdem ersetzen,
    # wenn genug Pixelfläche (tenant-spezifischer Header).
    if area >= 55_000:
        return True
    if longest >= 96:
        return True
    if kb >= 12 and area >= 2500:
        return True
    return False


def _fit_media(original_bytes: bytes, replacement_path: Path) -> bytes:
    """Skaliert Ersatz auf Original-Dimensionen; behält PNG vs JPEG bei."""
    orig = Image.open(io.BytesIO(original_bytes))
    tw, th = orig.size
    orig_fmt = (orig.format or "PNG").upper()

    repl = Image.open(replacement_path)
    repl = repl.convert("RGB") if repl.mode not in ("RGB", "L") else repl.convert("RGB")
    repl.thumbnail((tw, th), Image.LANCZOS)
    canvas = Image.new("RGB", (tw, th), (255, 255, 255))
    ox = (tw - repl.width) // 2
    oy = (th - repl.height) // 2
    canvas.paste(repl, (ox, oy))

    buf = io.BytesIO()
    if orig_fmt == "JPEG":
        canvas.save(buf, format="JPEG", quality=90, optimize=True)
    else:
        canvas.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


def opc_zip_deep_clean_and_media(
    docx_path: Path,
    *,
    explicit_media_map: dict[str, str],
    pool: list[Path],
) -> tuple[int, int, int]:
    """Ein Durchgang: OPC-XML Text + Medien tauschen.

    Returns ``(xml_hits, media_bulk, media_explicit)``
    """
    tmp = docx_path.with_suffix(".opc.tmp.docx")
    xml_hits = 0
    media_bulk = 0
    media_explicit = 0
    pool_i = 0

    explicit_paths = {k: HERE / v for k, v in explicit_media_map.items()}

    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                fn = info.filename

                # Alle OPC-XML und Relationship-Parts (Hyperlink-Ziele liegen in *.rels).
                if fn.endswith(".xml") or fn.endswith(".rels"):
                    text = data.decode("utf-8")
                    orig = text
                    for a, b in OPC_XML_TEXT_REPLACEMENTS:
                        if a in text:
                            text = text.replace(a, b)
                    if text != orig:
                        xml_hits += 1  # eine Datei mit mind. einem Treffer
                    data = text.encode("utf-8")

                elif fn.startswith("word/media/"):
                    name = Path(fn).name
                    raw_in = bytes(data)

                    if name in explicit_paths:
                        src = explicit_paths[name]
                        if not src.exists():
                            raise FileNotFoundError(f"image_mapping: fehlt {src}")
                        data = _fit_media(raw_in, src)
                        media_explicit += 1
                    elif should_replace_embedded_media(name, raw_in):
                        src = pool[pool_i % len(pool)]
                        pool_i += 1
                        data = _fit_media(raw_in, src)
                        media_bulk += 1

                zout.writestr(info, data)

    docx_path.unlink()
    tmp.rename(docx_path)
    return xml_hits, media_bulk, media_explicit


def verify_no_almirall_opc(docx_path: Path) -> None:
    bad: list[str] = []
    with zipfile.ZipFile(docx_path) as z:
        for name in z.namelist():
            if not (name.endswith(".xml") or name.endswith(".rels")):
                continue
            raw = z.read(name).decode("utf-8", errors="replace")
            if "almirall" in raw.lower():
                bad.append(name)
    if bad:
        raise RuntimeError(
            "QA fehlgeschlagen — noch 'almirall' in: " + ", ".join(sorted(set(bad)))
        )


# --- Strukturelle Transformation (python-docx), unveränderte Logik Kern ---


def _max_para_index(body) -> int:
    pi = -1
    for child in body:
        if child.tag == qn("w:p"):
            pi += 1
    return pi


def collect_para_slice(body, p_lo: int, p_hi: int) -> list:
    out: list = []
    pi = -1
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p"):
            pi += 1
            if p_lo <= pi <= p_hi:
                out.append(child)
        elif child.tag == qn("w:tbl"):
            if p_lo <= pi <= p_hi:
                out.append(child)
    return out


def remove_elems(body, elems: list) -> None:
    for el in elems:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def clone_para_with_text(proto_p_el, text: str) -> object:
    new_el = deepcopy(proto_p_el)
    for child in list(new_el):
        if child.tag == qn("w:r"):
            new_el.remove(child)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    new_el.append(r)
    return new_el


def capture_financial_prototypes(doc: Document):
    body = doc.element.body
    pi = -1
    proto_h1 = proto_h2 = proto_li = None
    for child in body:
        if child.tag != qn("w:p"):
            continue
        pi += 1
        if pi == 155:
            proto_h1 = deepcopy(child)
        elif pi == 157:
            proto_li = deepcopy(child)
        elif pi == 160:
            proto_h2 = deepcopy(child)
    if proto_h1 is None or proto_h2 is None or proto_li is None:
        raise RuntimeError("Konnte SPA-Prototyp-Absätze (155/157/160) nicht finden.")
    return proto_h1, proto_h2, proto_li


def build_financials_elements(proto_h1, proto_h2, proto_li) -> list:
    elems: list = []

    def h1(txt: str):
        elems.append(clone_para_with_text(deepcopy(proto_h1), txt))

    def h2(txt: str):
        elems.append(clone_para_with_text(deepcopy(proto_h2), txt))

    def li(txt: str):
        elems.append(clone_para_with_text(deepcopy(proto_li), txt))

    h1("PROJECT FINANCIALS")
    li(
        "Financial planning completes the Planner triangle together with Attributes and "
        "the Single Project Allocation grid. Hourly / Advanced rates are administered by Finance; "
        "PM/SM maintain Cost Plan overrides where permitted and interpret Budget vs Actual.",
    )
    h2("Rates — Hourly & Advanced")
    li(
        "Hourly Rate — blended €/h per named or generic resource for a validity window; "
        "shown as read-only context while planning.",
    )
    li(
        "Advanced Rate — dated brackets or uplifts (e.g. vendors). Tempus resolves the effective €/h "
        "per period; reconcile surprises via the SPA grid Cost view.",
    )
    h2("Cost Plan & Budget vs Actual")
    li(
        "Cost Plan — derived from Hours × effective rate; Finance-approved manual cells capture "
        "non-hour spend (licences, PO-driven costs).",
    )
    li(
        "Budget vs Actual — compares planned cost with booked hours × rate. Align granularity "
        "(month / quarter / year) with Knauf steering cadence.",
    )
    h2("Lock Periods & Re-forecast")
    li(
        "Lock Periods — closed financial buckets prevent silent retro-edits; escalate hard-lock "
        "changes through PMO / Finance.",
    )
    li(
        "Re-forecast — refresh forward buckets after actuals land; prefer SPA edits first so "
        "mandays / euros remain coherent before overriding Cost Plan cells.",
    )
    return elems


def reorder_and_trim(doc: Document, protos: tuple) -> None:
    body = doc.element.body
    p_max = _max_para_index(body)

    vm = collect_para_slice(body, 121, 137)
    attr = collect_para_slice(body, 138, 154)
    spa = collect_para_slice(body, 155, 239)
    bpa = collect_para_slice(body, 240, p_max)

    remove_elems(body, vm + attr + spa + bpa)

    financial_elems = build_financials_elements(*protos)

    for el in attr + spa + financial_elems + vm:
        body.append(el)


def apply_text_replacements(doc: Document) -> int:
    reps = [
        ("https://almirall.tempus-resource.com/sg/", KNAUF_TEMPUS_URL),
        ("almirall.tempus-resource.com", "knaufdev.tempus-resource.com"),
        ("Almirall Hermal", "Knauf SE"),
        ("Almirall", "Knauf"),
        ("almirall", "knauf"),
        ("Continue With Microsoft", "Knauf SSO"),
        ("Project Management GRId", "Project Management Grid"),
        ("Month is the recommended granularity.", "Month is the recommended granularity for Knauf."),
        (
            "Hours is the recommended unit for capacity planning.",
            "Hours are recommended for Knauf capacity planning (switch to Mandays when steering leadership).",
        ),
        ("Sönnke Grossmann", "another named resource"),
        ("Sönnke Grossman", "another named resource"),
        ("QP Generic (SAB)", "Demand-planning generics"),
        (
            "All PM users will have access to 6 tiles on their homepage.",
            "Planner roles (PM/SM and blended PM/SM + PPM) see the homepage tiles provisioned for them — typically "
            "Project Management and Report Management.",
        ),
        (
            "\t\t*Areas covered in this guide are highlighted in yellow.",
            "\t\t*Highlighted screenshots mirror Knauf Planner onboarding (PM/SM). RM-only surfaces such as Bulk Project Allocation "
            "are intentionally omitted from this edition.",
        ),
    ]
    n = 0

    def repl_block(block):
        nonlocal n
        for para in block:
            full = para.text
            if not full:
                continue
            new_full = full
            for a, b in reps:
                new_full = new_full.replace(a, b)
            if new_full != full:
                if para.runs:
                    para.runs[0].text = new_full
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.add_run(new_full)
                n += 1

    repl_block(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                repl_block(cell.paragraphs)
    for sec in doc.sections:
        repl_block(sec.header.paragraphs)
        repl_block(sec.footer.paragraphs)
    return n


def patch_homescreen_table(doc: Document) -> None:
    tbl = doc.tables[0]
    while len(tbl.rows) > 1:
        tbl._tbl.remove(tbl.rows[-1]._tr)
    rows_desc = [
        (
            "Project Management",
            "View project repository | Create projects with / without template | "
            "Open projects | Maintain allocations on the SPA grid",
        ),
        (
            "Report Management",
            "Consume shared reports | Export to Excel / PDF where licensed",
        ),
    ]
    for name, desc in rows_desc:
        tbl.add_row()
        row = tbl.rows[-1]
        row.cells[0].text = name
        row.cells[1].text = desc


def insert_scope_note_after_overview(doc: Document) -> None:
    body = doc.element.body
    pi = -1
    overview_el = None
    proto_el = None
    for child in body:
        if child.tag != qn("w:p"):
            continue
        pi += 1
        para = DocxParagraph(child, doc)
        if para.style.name.startswith("Heading") and (para.text or "").strip() == "Overview":
            overview_el = child
            nxt = overview_el.getnext()
            while nxt is not None and nxt.tag != qn("w:p"):
                nxt = nxt.getnext()
            proto_el = nxt
            break
    if overview_el is None or proto_el is None:
        return
    note_el = clone_para_with_text(
        proto_el,
        "Planner scope (Knauf Phase 1): PM/SM and blended PM/SM + PPM are in scope. Resource Manager chapters "
        "(Bulk Project Allocation Flatgrid and enterprise-wide replace flows) ship separately.",
    )
    overview_el.addnext(note_el)


def main() -> None:
    if not TEMPLATE.exists():
        raise SystemExit(f"Template fehlt: {TEMPLATE}")

    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    SCREENSHOTS_LOCAL.mkdir(parents=True, exist_ok=True)

    pool = discover_knauf_screenshot_pool()

    shutil.copy(TEMPLATE, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)

    protos = capture_financial_prototypes(doc)
    reorder_and_trim(doc, protos)

    n_para_runs = apply_text_replacements(doc)
    patch_homescreen_table(doc)
    insert_scope_note_after_overview(doc)

    img_map_path = CONFIG_DIR / "image_mapping.json"
    explicit: dict[str, str] = {}
    if img_map_path.exists():
        explicit = json.loads(img_map_path.read_text(encoding="utf-8")).get(
            "replacements", {}
        )

    doc.save(OUTPUT_DOCX)

    xh, mb, me = opc_zip_deep_clean_and_media(
        OUTPUT_DOCX,
        explicit_media_map=explicit,
        pool=pool,
    )

    verify_no_almirall_opc(OUTPUT_DOCX)

    print(f"python-docx Absatz-/Tabellen-Ersetzungen (Run-Batches): {n_para_runs}")
    print(f"OPC XML-Dateien mit Textänderungen: {xh}")
    print(f"Medien automatisch ersetzt (Pool): {mb}")
    print(f"Medien per image_mapping.json: {me}")
    print(f"Ausgabe: {OUTPUT_DOCX}")
    print("QA: Kein 'almirall' mehr in allen *.xml / *.rels des Pakets ✓")


if __name__ == "__main__":
    main()
