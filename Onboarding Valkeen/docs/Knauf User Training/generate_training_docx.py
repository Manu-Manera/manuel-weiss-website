"""Build the Knauf planner (PM/SM & RM) User Training DOCX.

Reads markdown files from ``content/`` in the order defined by ``SECTIONS``,
resolves image references against the training folder root, and renders the
final ``Knauf_User_Training_Project_Manager.docx``.

Markdown features supported (intentionally a small subset):

- ``# Heading 1`` / ``## Heading 2`` / ``### Heading 3``
- Paragraphs
- Unordered lists ``- item`` and ordered lists ``1. item``
- Inline emphasis with ``**bold**`` and ``*italic*``
- Inline code with backticks
- Inline links (rendered as plain text)
- Image lines: ``![Caption](screenshots/file.png)``
- Callouts: a paragraph starting with ``> Note:`` becomes a note callout

The output is page-number-stamped and carries the Valkeen footer block.
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


HERE = Path(__file__).resolve().parent
CONTENT_DIR = HERE / "content"
SCREENSHOT_DIR = HERE / "screenshots"
OUTPUT_FILE = HERE / "Knauf_User_Training_Project_Manager.docx"

SECTIONS: list[str] = [
    "00_cover.md",
    "01_contents.md",
    "02_overview.md",
    "03_shortcuts.md",
    "04_navigation.md",
    "05_create_project.md",
    "06_master_data.md",
    "07_project_vs_service.md",
    "08_resource_planning.md",
    "09_cost_planning.md",
    "10_bpa_flatgrid.md",
    "11_reporting.md",
]

VALKEEN_BLUE = RGBColor(0x0B, 0x3D, 0x91)
TEXT_GRAY = RGBColor(0x3A, 0x3A, 0x3A)
NOTE_BG = "EAF1FB"

PAGE_WIDTH_CM = 21.0
PAGE_MARGIN_CM = 2.0
USABLE_WIDTH_CM = PAGE_WIDTH_CM - 2 * PAGE_MARGIN_CM


@dataclass
class Block:
    kind: str
    text: str = ""
    items: list[str] | None = None
    ordered: bool = False
    image_path: str = ""
    image_caption: str = ""


def parse_markdown(md_text: str) -> list[Block]:
    """Parse a markdown string into a flat list of typed blocks."""
    blocks: list[Block] = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        m_img = re.match(r"^!\[(?P<cap>[^\]]*)\]\((?P<path>[^)]+)\)\s*$", line)
        if m_img:
            blocks.append(Block(
                kind="image",
                image_path=m_img.group("path"),
                image_caption=m_img.group("cap"),
            ))
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(Block(kind="h3", text=line[4:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(Block(kind="h2", text=line[3:].strip()))
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(Block(kind="h1", text=line[2:].strip()))
            i += 1
            continue
        if line.startswith("> "):
            blocks.append(Block(kind="note", text=line[2:].strip()))
            i += 1
            continue
        if re.match(r"^[-*]\s+", line):
            items: list[str] = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].rstrip() or ""):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].rstrip()))
                i += 1
            blocks.append(Block(kind="list", items=items, ordered=False))
            continue
        if re.match(r"^\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].rstrip() or ""):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].rstrip()))
                i += 1
            blocks.append(Block(kind="list", items=items, ordered=True))
            continue
        paragraph_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not _starts_block(lines[i]):
            paragraph_lines.append(lines[i].rstrip())
            i += 1
        blocks.append(Block(kind="p", text=" ".join(paragraph_lines)))
    return blocks


def _starts_block(line: str) -> bool:
    line = line.lstrip()
    return (
        line.startswith("#")
        or line.startswith("> ")
        or line.startswith("- ")
        or line.startswith("* ")
        or re.match(r"^\d+\.\s+", line) is not None
        or re.match(r"^!\[.*\]\(.+\)\s*$", line) is not None
    )


INLINE_PATTERN = re.compile(
    r"(\*\*(?P<bold>[^*]+)\*\*|"
    r"\*(?P<ital>[^*]+)\*|"
    r"`(?P<code>[^`]+)`|"
    r"\[(?P<linkt>[^\]]+)\]\((?P<linku>[^)]+)\))"
)


def add_inline(paragraph, text: str) -> None:
    """Add a text string to a paragraph, applying inline formatting."""
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        if match.group("bold"):
            run = paragraph.add_run(match.group("bold"))
            run.bold = True
        elif match.group("ital"):
            run = paragraph.add_run(match.group("ital"))
            run.italic = True
        elif match.group("code"):
            run = paragraph.add_run(match.group("code"))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif match.group("linkt"):
            run = paragraph.add_run(match.group("linkt"))
            run.font.color.rgb = VALKEEN_BLUE
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(document: Document) -> None:
    """Set default font and tune the built-in heading styles."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_GRAY

    for level, size in [(1, 22), (2, 16), (3, 13)]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = VALKEEN_BLUE


def configure_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(PAGE_MARGIN_CM)
    section.bottom_margin = Cm(PAGE_MARGIN_CM)
    section.left_margin = Cm(PAGE_MARGIN_CM)
    section.right_margin = Cm(PAGE_MARGIN_CM)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = ""

    line1 = footer.paragraphs[0]
    line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line1.add_run(
        "Valkeen GmbH \u2219 Bahnhofstrasse 64 \u2219 8001 Zurich, Switzerland"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_GRAY

    line2 = footer.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line2.add_run("www.valkeen.com \u2219 \u00a9 2026 Valkeen GmbH. All rights reserved.")
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_GRAY

    line3 = footer.add_paragraph()
    line3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line3.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_GRAY
    add_page_number_field(line3)


def render_blocks(document: Document, blocks: Iterable[Block], section_index: int) -> None:
    for block in blocks:
        if block.kind == "h1":
            if section_index > 0:
                p = document.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
            heading = document.add_heading(level=1)
            add_inline(heading, block.text)
        elif block.kind == "h2":
            heading = document.add_heading(level=2)
            add_inline(heading, block.text)
        elif block.kind == "h3":
            heading = document.add_heading(level=3)
            add_inline(heading, block.text)
        elif block.kind == "p":
            paragraph = document.add_paragraph()
            add_inline(paragraph, block.text)
        elif block.kind == "list":
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items or []:
                paragraph = document.add_paragraph(style=style)
                add_inline(paragraph, item)
        elif block.kind == "note":
            table = document.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, NOTE_BG)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run("Note: ")
            run.bold = True
            run.font.color.rgb = VALKEEN_BLUE
            add_inline(paragraph, block.text)
            document.add_paragraph()
        elif block.kind == "image":
            full_path = (HERE / block.image_path).resolve()
            if not full_path.exists():
                placeholder = document.add_paragraph()
                run = placeholder.add_run(
                    f"[Missing screenshot: {block.image_path}]"
                )
                run.italic = True
                run.font.color.rgb = RGBColor(0xB0, 0x0B, 0x1E)
                continue
            with Image.open(full_path) as img:
                width_px, _ = img.size
            dpi = 96
            width_cm_native = width_px / dpi * 2.54
            width_cm = min(width_cm_native, USABLE_WIDTH_CM)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(full_path), width=Cm(width_cm))
            if block.image_caption:
                caption = document.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(block.image_caption)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = TEXT_GRAY


def build() -> Path:
    if not CONTENT_DIR.exists():
        raise SystemExit(f"Content directory missing: {CONTENT_DIR}")
    document = Document()
    configure_page_layout(document)
    configure_styles(document)

    for index, filename in enumerate(SECTIONS):
        md_path = CONTENT_DIR / filename
        if not md_path.exists():
            print(f"[warn] skipping missing content file: {filename}", file=sys.stderr)
            continue
        blocks = parse_markdown(md_path.read_text(encoding="utf-8"))
        render_blocks(document, blocks, section_index=index)

    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    output = build()
    print(f"Wrote: {output}")
