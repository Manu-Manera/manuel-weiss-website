"""Build the Knauf Tempus Resource Quick Reference Guide (Planners PM/SM & RM).

Reads markdown files from ``content/`` in the order defined by ``SECTIONS`` and
renders ``Knauf_Tempus_Quick_Reference_Guide_PM-RM.docx``.

Markdown features supported (kept minimal and deterministic):

- ``# H1`` / ``## H2`` / ``### H3``
- Paragraphs, bullet lists (``- item``), ordered lists (``1. item``)
- Inline ``**bold**``, ``*italic*``, ``` `code` ``` and ``[link](url)``
- Images: ``![Caption](screenshots/file.png)``
- Note callout: a paragraph starting with ``> Note:``
- Tip / Warning callouts: paragraphs starting with ``> Tip:`` or ``> Warning:``
- Quick Reference extensions (custom):
    - ``:::roles PM/SM | RM`` (single line)   - role badges next to the title
    - ``:::aag`` ... ``:::``                  - "At a glance" highlight box
    - ``:::steps`` ... ``:::``                - numbered "How to" box
    - ``:::toc``                              - inserts a live Word TOC field
- Pipe tables:
    ``| Header 1 | Header 2 |``
    ``| --- | --- |``
    ``| cell | cell |``
- Inline images / icons (rendered inline in a paragraph):
    ``Click ![edit|h=14](lib:icons/icon_edit.png) to edit.``
    ``Click :icon[edit] to edit.``    (resolves via _tempus-assets/index.json)
- Block images may use the same ``|h=NN`` (pt) suffix to override height, e.g.
    ``![Caption|h=180](lib:screens/screen_homescreen_pmsm.png)``
- Library paths starting with ``lib:`` are resolved against the shared
  ``_tempus-assets`` directory (one level above the per-doc folder).
"""

from __future__ import annotations

import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


HERE = Path(__file__).resolve().parent
CONTENT_DIR = HERE / "content"
SCREENSHOT_DIR = HERE / "screenshots"
LIBRARY_DIR = (HERE.parent / "_tempus-assets").resolve()
OUTPUT_FILE = HERE / "Knauf_Tempus_QRG_PM-SM.docx"


SECTIONS: list[str] = [
    "00_cover.md",
    "01_contents.md",
    "02_intro.md",
    "03_sign_in_and_homescreen.md",
    "04_create_project.md",
    "05_setup_attributes.md",
    "06_setup_resources.md",
    "07_setup_financials.md",
    "08_find_and_open.md",
    "09_appendix.md",
]


VALKEEN_BLUE = RGBColor(0x0B, 0x3D, 0x91)
TEXT_GRAY = RGBColor(0x3A, 0x3A, 0x3A)
NOTE_BG = "EAF1FB"
AAG_BG = "EAF3FF"
STEPS_BG = "F5F0FF"
TIP_BG = "EAF7EE"
WARN_BG = "FBEAEA"
TABLE_HEADER_BG = "0B3D91"
TABLE_HEADER_TEXT = "FFFFFF"
TABLE_ROW_ALT_BG = "F5F7FB"

BADGE_COLORS = {
    "PM/SM": "0B3D91",
    "RM": "1E7A3C",
    "Both": "5B2A86",
    "Controlling": "B58105",
    "PPM": "1B6F8E",
    "Admin": "B0312F",
}
BADGE_DEFAULT = "5A5A5A"

PAGE_WIDTH_CM = 21.0
PAGE_MARGIN_CM = 2.0
USABLE_WIDTH_CM = PAGE_WIDTH_CM - 2 * PAGE_MARGIN_CM


@dataclass
class Block:
    kind: str
    text: str = ""
    items: list[str] = field(default_factory=list)
    ordered: bool = False
    image_path: str = ""
    image_caption: str = ""
    image_height_pt: int | None = None
    rows: list[list[str]] = field(default_factory=list)
    note_kind: str = "note"


# ----------------------------- Markdown parser ----------------------------- #


_FENCE_OPEN = re.compile(r"^:::(?P<name>roles|aag|steps|toc|highlight)(?:\s+(?P<arg>.+))?\s*$")
_FENCE_CLOSE = re.compile(r"^:::\s*$")
_IMAGE = re.compile(r"^!\[(?P<cap>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")
_NOTE_PREFIX = re.compile(r"^>\s*(?P<kind>Note|Tip|Warning):\s*(?P<text>.*)$")
_NOTE_PLAIN = re.compile(r"^>\s+(?P<text>.+)$")
_LIST_BULLET = re.compile(r"^[-*]\s+(.*)$")
_LIST_ORDERED = re.compile(r"^\d+\.\s+(.*)$")
_TABLE_ROW = re.compile(r"^\|(?P<body>.+)\|\s*$")
_TABLE_SEP = re.compile(r"^\|[\s:|-]+\|\s*$")
# Inline tokens for ![alt|h=14](path) and :icon[name]
_INLINE_IMAGE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")
_INLINE_ICON = re.compile(r":icon\[(?P<name>[^\]]+)\]")


# ----------------------------- Asset library ------------------------------ #


_LIB_INDEX_CACHE: dict | None = None


def _load_library_index() -> dict:
    global _LIB_INDEX_CACHE
    if _LIB_INDEX_CACHE is None:
        idx_path = LIBRARY_DIR / "index.json"
        if idx_path.exists():
            try:
                data = json.loads(idx_path.read_text(encoding="utf-8"))
                _LIB_INDEX_CACHE = data.get("assets", {}) or {}
            except Exception as exc:  # pragma: no cover
                print(f"[warn] could not read library index: {exc}", file=sys.stderr)
                _LIB_INDEX_CACHE = {}
        else:
            _LIB_INDEX_CACHE = {}
    return _LIB_INDEX_CACHE


def resolve_image_path(path_token: str) -> Path:
    """Resolve a Markdown image path. ``lib:foo/bar.png`` -> central library."""
    if path_token.startswith("lib:"):
        return (LIBRARY_DIR / path_token[4:]).resolve()
    return (HERE / path_token).resolve()


def resolve_icon(name: str) -> tuple[Path | None, int]:
    """Look up an icon by id in the library index. Returns (path, height_pt)."""
    asset = _load_library_index().get(name)
    if not asset:
        return None, 12
    file = asset.get("file", "")
    height = int(asset.get("default_inline_height_pt", 12))
    return (LIBRARY_DIR / file).resolve(), height


def parse_image_token(path_token: str) -> tuple[str, int | None]:
    """Split a Markdown image alt-text/path that may carry a ``|h=NN`` suffix.

    Accepts the suffix on either the alt portion or the path portion. Returns
    ``(cleaned_token, height_pt_or_None)``.
    """
    height: int | None = None
    cleaned = path_token
    m = re.search(r"\|h=(\d+)", cleaned)
    if m:
        try:
            height = int(m.group(1))
        except ValueError:
            height = None
        cleaned = (cleaned[: m.start()] + cleaned[m.end():]).strip()
    return cleaned, height


def parse_markdown(md_text: str) -> list[Block]:
    blocks: list[Block] = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if not line.strip():
            i += 1
            continue

        # Stray closing fence outside any open block – treat as a no-op.
        if _FENCE_CLOSE.match(line):
            i += 1
            continue

        # Fenced custom blocks
        m_fence = _FENCE_OPEN.match(line)
        if m_fence:
            name = m_fence.group("name")
            arg = m_fence.group("arg") or ""
            if name == "toc":
                blocks.append(Block(kind="toc"))
                i += 1
                continue
            if name == "roles":
                roles = [r.strip() for r in re.split(r"\||,", arg) if r.strip()]
                blocks.append(Block(kind="roles", items=roles))
                i += 1
                continue
            if name == "highlight":
                # ``:::highlight image=lib:... caption="..."`` – args on opening line,
                # body optional and ignored.
                attrs = _parse_kv_args(arg)
                # Consume optional body up to closing fence
                i += 1
                while i < len(lines) and not _FENCE_CLOSE.match(lines[i].rstrip()):
                    i += 1
                if i < len(lines) and _FENCE_CLOSE.match(lines[i].rstrip()):
                    i += 1
                path_raw = attrs.get("image", "")
                cleaned_path, height = parse_image_token(path_raw)
                blocks.append(
                    Block(
                        kind="highlight",
                        image_path=cleaned_path,
                        image_caption=attrs.get("caption", ""),
                        image_height_pt=height,
                    )
                )
                continue
            # aag / steps – consume until closing fence
            i += 1
            body_items: list[str] = []
            while i < len(lines) and not _FENCE_CLOSE.match(lines[i].rstrip()):
                stripped = lines[i].strip()
                if not stripped:
                    i += 1
                    continue
                m_b = _LIST_BULLET.match(stripped)
                m_o = _LIST_ORDERED.match(stripped)
                if m_b:
                    body_items.append(m_b.group(1))
                elif m_o:
                    body_items.append(m_o.group(1))
                else:
                    body_items.append(stripped)
                i += 1
            if i < len(lines) and _FENCE_CLOSE.match(lines[i].rstrip()):
                i += 1
            blocks.append(Block(kind=name, items=body_items))
            continue

        # Image (block, on its own line)
        m_img = _IMAGE.match(line)
        if m_img:
            raw_path = m_img.group("path")
            raw_cap = m_img.group("cap")
            cleaned_path, h_from_path = parse_image_token(raw_path)
            cleaned_cap, h_from_cap = parse_image_token(raw_cap)
            height = h_from_path if h_from_path is not None else h_from_cap
            blocks.append(
                Block(
                    kind="image",
                    image_path=cleaned_path,
                    image_caption=cleaned_cap,
                    image_height_pt=height,
                )
            )
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            blocks.append(Block(kind="h3", text=line[4:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(Block(kind="h2", text=line[3:].strip()))
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(Block(kind="h1", text=line[2:].strip()))
            i += 1
            continue

        # Note / Tip / Warning — consume *all* consecutive ``>`` lines so
        # a multi-line blockquote renders as a single note banner rather
        # than a stack of one-line banners.
        if line.startswith(">"):
            note_lines: list[str] = []
            kind = "note"
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                stripped = re.sub(r"^>\s?", "", lines[i].rstrip())
                if not note_lines:
                    m_typed = re.match(r"^(?P<kind>Note|Tip|Warning):\s*(?P<text>.*)$", stripped)
                    if m_typed:
                        kind = m_typed.group("kind").lower()
                        stripped = m_typed.group("text")
                note_lines.append(stripped.strip())
                i += 1
            text = " ".join(s for s in note_lines if s)
            blocks.append(Block(kind="note", note_kind=kind, text=text))
            continue

        # Tables (pipe)
        if _TABLE_ROW.match(line) and (i + 1) < len(lines) and _TABLE_SEP.match(lines[i + 1].rstrip()):
            header = _split_pipe_row(line)
            i += 2
            body_rows: list[list[str]] = []
            while i < len(lines) and _TABLE_ROW.match(lines[i].rstrip()):
                body_rows.append(_split_pipe_row(lines[i].rstrip()))
                i += 1
            blocks.append(Block(kind="table", items=header, rows=body_rows))
            continue

        # Bullet list — also supports multi-line items
        if _LIST_BULLET.match(line):
            items: list[str] = []
            while i < len(lines):
                cur = lines[i].rstrip()
                m = _LIST_BULLET.match(cur or "")
                if m:
                    items.append(m.group(1))
                    i += 1
                    continue
                if cur.startswith("  ") and items:
                    items[-1] += " " + cur.strip()
                    i += 1
                    continue
                break
            blocks.append(Block(kind="list", items=items, ordered=False))
            continue

        # Ordered list — supports multi-line items (continuation lines
        # are indented with at least two spaces).
        if _LIST_ORDERED.match(line):
            items: list[str] = []
            while i < len(lines):
                cur = lines[i].rstrip()
                m = _LIST_ORDERED.match(cur or "")
                if m:
                    items.append(m.group(1))
                    i += 1
                    continue
                if cur.startswith("  ") and items:
                    items[-1] += " " + cur.strip()
                    i += 1
                    continue
                break
            blocks.append(Block(kind="list", items=items, ordered=True))
            continue

        # Paragraph (collect continuation lines)
        paragraph_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not _starts_block(lines[i]):
            paragraph_lines.append(lines[i].rstrip())
            i += 1
        blocks.append(Block(kind="p", text=" ".join(paragraph_lines)))

    return blocks


def _split_pipe_row(line: str) -> list[str]:
    body = line.strip().strip("|")
    return [cell.strip() for cell in body.split("|")]


def _parse_kv_args(arg: str) -> dict[str, str]:
    """Parse ``key=value key2="quoted value"`` argument strings."""
    result: dict[str, str] = {}
    pos = 0
    while pos < len(arg):
        while pos < len(arg) and arg[pos].isspace():
            pos += 1
        if pos >= len(arg):
            break
        eq = arg.find("=", pos)
        if eq < 0:
            break
        key = arg[pos:eq].strip()
        pos = eq + 1
        if pos < len(arg) and arg[pos] in ('"', "'"):
            quote = arg[pos]
            pos += 1
            end = arg.find(quote, pos)
            if end < 0:
                end = len(arg)
            value = arg[pos:end]
            pos = end + 1
        else:
            end = pos
            while end < len(arg) and not arg[end].isspace():
                end += 1
            value = arg[pos:end]
            pos = end
        if key:
            result[key] = value
    return result


def _starts_block(line: str) -> bool:
    s = line.lstrip()
    return (
        s.startswith("#")
        or s.startswith(">")
        or s.startswith("- ")
        or s.startswith("* ")
        or _LIST_ORDERED.match(s) is not None
        or _IMAGE.match(s) is not None
        or _TABLE_ROW.match(s) is not None
        or _FENCE_OPEN.match(s) is not None
        or _FENCE_CLOSE.match(s) is not None
    )


# ----------------------------- Inline formatting --------------------------- #


INLINE_PATTERN = re.compile(
    r"(!\[(?P<imgalt>[^\]]*)\]\((?P<imgpath>[^)]+)\)|"
    r":icon\[(?P<iconname>[^\]]+)\]|"
    r"\*\*(?P<bold>[^*]+)\*\*|"
    r"\*(?P<ital>[^*]+)\*|"
    r"`(?P<code>[^`]+)`|"
    r"\[(?P<linkt>[^\]]+)\]\((?P<linku>[^)]+)\))"
)


DEFAULT_INLINE_ICON_HEIGHT_PT = 12


def _add_inline_image(paragraph, path: Path, height_pt: int) -> None:
    """Insert an inline image into the run sequence of ``paragraph``."""
    if not path.exists():
        run = paragraph.add_run(f"[missing:{path.name}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xB0, 0x0B, 0x1E)
        return
    run = paragraph.add_run()
    try:
        run.add_picture(str(path), height=Pt(height_pt))
    except Exception as exc:  # pragma: no cover
        run.text = f"[img-error:{path.name}:{exc}]"
        run.font.color.rgb = RGBColor(0xB0, 0x0B, 0x1E)


def add_inline(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        if match.group("imgpath") is not None:
            raw_path = match.group("imgpath")
            raw_alt = match.group("imgalt") or ""
            cleaned_path, h_path = parse_image_token(raw_path)
            _, h_alt = parse_image_token(raw_alt)
            height = h_path if h_path is not None else h_alt
            target = resolve_image_path(cleaned_path)
            _add_inline_image(
                paragraph,
                target,
                height or DEFAULT_INLINE_ICON_HEIGHT_PT,
            )
        elif match.group("iconname"):
            name = match.group("iconname").strip()
            target, height_pt = resolve_icon(name)
            if target is None:
                run = paragraph.add_run(f"[icon?{name}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0xB0, 0x0B, 0x1E)
            else:
                _add_inline_image(paragraph, target, height_pt)
        elif match.group("bold"):
            run = paragraph.add_run(match.group("bold"))
            run.bold = True
        elif match.group("ital"):
            run = paragraph.add_run(match.group("ital"))
            run.italic = True
        elif match.group("code"):
            run = paragraph.add_run(match.group("code"))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif match.group("linkt"):
            run = paragraph.add_run(match.group("linkt"))
            run.font.color.rgb = VALKEEN_BLUE
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ----------------------------- Low-level helpers --------------------------- #


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex: str = "C8D2E8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc_field(document: Document) -> None:
    """Insert a real Word TOC field. Word will populate on open / F9."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:r")
    fallback_text = OxmlElement("w:t")
    fallback_text.text = (
        "Right-click here and choose 'Update Field' (or press F9) "
        "to build the table of contents."
    )
    fallback.append(fallback_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fallback)
    run._r.append(fld_end)


# ----------------------------- Styles & layout ----------------------------- #


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_GRAY

    for level, size in [(1, 22), (2, 16), (3, 13)]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = VALKEEN_BLUE


def configure_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(PAGE_MARGIN_CM)
    section.bottom_margin = Cm(PAGE_MARGIN_CM)
    section.left_margin = Cm(PAGE_MARGIN_CM)
    section.right_margin = Cm(PAGE_MARGIN_CM)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = ""

    line1 = footer.paragraphs[0]
    line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line1.add_run(
        "Valkeen GmbH \u2219 Bahnhofstrasse 64 \u2219 8001 Zurich, Switzerland"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_GRAY

    line2 = footer.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line2.add_run(
        "www.valkeen.com \u2219 \u00a9 2026 Valkeen GmbH. All rights reserved."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_GRAY

    line3 = footer.add_paragraph()
    line3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line3.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_GRAY
    add_page_number_field(line3)


# ----------------------------- Block renderers ----------------------------- #


def render_roles(document: Document, roles: list[str]) -> None:
    if not roles:
        return
    table = document.add_table(rows=1, cols=len(roles))
    table.autofit = True
    for idx, role in enumerate(roles):
        cell = table.cell(0, idx)
        color = BADGE_COLORS.get(role, BADGE_DEFAULT)
        set_cell_shading(cell, color)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(role)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_table_layout_fixed(table, total_cm=min(USABLE_WIDTH_CM, 2.6 * len(roles)))
    document.add_paragraph()


def render_box(document: Document, title: str, items: list[str], bg_hex: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_hex)
    set_cell_borders(cell, color_hex="C8D2E8")
    para = cell.paragraphs[0]
    run = para.add_run(title)
    run.bold = True
    run.font.color.rgb = VALKEEN_BLUE
    run.font.size = Pt(11)
    for idx, item in enumerate(items):
        line = cell.add_paragraph()
        prefix = f"{idx + 1}. " if title.lower().startswith("how to") else "\u2022 "
        run = line.add_run(prefix)
        run.bold = True
        run.font.color.rgb = VALKEEN_BLUE
        add_inline(line, item)
    document.add_paragraph()


def render_table(document: Document, header: list[str], rows: list[list[str]]) -> None:
    if not header:
        return
    cols = len(header)
    docx_table = document.add_table(rows=1 + len(rows), cols=cols)
    docx_table.autofit = True

    # Header row
    for c, txt in enumerate(header):
        cell = docx_table.cell(0, c)
        set_cell_shading(cell, TABLE_HEADER_BG)
        para = cell.paragraphs[0]
        run = para.add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Body rows
    for r, row in enumerate(rows):
        for c in range(cols):
            value = row[c] if c < len(row) else ""
            cell = docx_table.cell(1 + r, c)
            if r % 2 == 1:
                set_cell_shading(cell, TABLE_ROW_ALT_BG)
            para = cell.paragraphs[0]
            add_inline(para, value)
            for run in para.runs:
                run.font.size = Pt(10)

    _set_table_layout_fixed(docx_table, total_cm=USABLE_WIDTH_CM)
    document.add_paragraph()


def render_note(document: Document, kind: str, text: str) -> None:
    label = {"note": "Note", "tip": "Tip", "warning": "Warning"}.get(kind, "Note")
    bg = {"note": NOTE_BG, "tip": TIP_BG, "warning": WARN_BG}.get(kind, NOTE_BG)
    color = {
        "note": VALKEEN_BLUE,
        "tip": RGBColor(0x1E, 0x7A, 0x3C),
        "warning": RGBColor(0xB0, 0x31, 0x2F),
    }.get(kind, VALKEEN_BLUE)
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg)
    para = cell.paragraphs[0]
    para.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = para.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = color
    add_inline(para, text)
    document.add_paragraph()


def render_image_block(
    document: Document,
    path_token: str,
    caption: str,
    height_pt: int | None = None,
) -> None:
    full_path = resolve_image_path(path_token)
    if not full_path.exists():
        placeholder = document.add_paragraph()
        run = placeholder.add_run(f"[Missing screenshot: {path_token}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xB0, 0x0B, 0x1E)
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    if height_pt:
        run.add_picture(str(full_path), height=Pt(height_pt))
    else:
        with Image.open(full_path) as img:
            width_px, _ = img.size
        dpi = 96
        width_cm_native = width_px / dpi * 2.54
        width_cm = min(width_cm_native, USABLE_WIDTH_CM)
        run.add_picture(str(full_path), width=Cm(width_cm))
    if caption:
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cap.add_run(caption)
        crun.italic = True
        crun.font.size = Pt(9)
        crun.font.color.rgb = TEXT_GRAY


def render_highlight_block(
    document: Document,
    path_token: str,
    caption: str,
    height_pt: int | None = None,
) -> None:
    """Highlight: image rendered inside a shaded 1x1 table with a yellow accent border."""
    full_path = resolve_image_path(path_token)
    if not full_path.exists():
        placeholder = document.add_paragraph()
        run = placeholder.add_run(f"[Missing highlight: {path_token}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xB0, 0x0B, 0x1E)
        return
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF8DC")
    set_cell_borders(cell, color_hex="F2B807")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    if height_pt:
        run.add_picture(str(full_path), height=Pt(height_pt))
    else:
        with Image.open(full_path) as img:
            width_px, _ = img.size
        dpi = 96
        width_cm_native = width_px / dpi * 2.54
        width_cm = min(width_cm_native, USABLE_WIDTH_CM - 1.0)
        run.add_picture(str(full_path), width=Cm(width_cm))
    if caption:
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cap.add_run(caption)
        crun.italic = True
        crun.font.size = Pt(9)
        crun.font.color.rgb = TEXT_GRAY
    document.add_paragraph()


def _set_table_layout_fixed(table, total_cm: float) -> None:
    cols = len(table.columns)
    if cols == 0:
        return
    width_per_col = Cm(total_cm / cols)
    for row in table.rows:
        for cell in row.cells:
            cell.width = width_per_col


def render_blocks(document: Document, blocks: Iterable[Block], section_index: int) -> None:
    for block in blocks:
        if block.kind == "h1":
            if section_index > 0:
                p = document.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
            heading = document.add_heading(level=1)
            add_inline(heading, block.text)
        elif block.kind == "h2":
            heading = document.add_heading(level=2)
            add_inline(heading, block.text)
        elif block.kind == "h3":
            heading = document.add_heading(level=3)
            add_inline(heading, block.text)
        elif block.kind == "p":
            paragraph = document.add_paragraph()
            add_inline(paragraph, block.text)
        elif block.kind == "list":
            if block.ordered:
                for index, item in enumerate(block.items, start=1):
                    paragraph = document.add_paragraph()
                    paragraph.paragraph_format.left_indent = Pt(18)
                    paragraph.paragraph_format.first_line_indent = Pt(-18)
                    run = paragraph.add_run(f"{index}.\t")
                    run.bold = True
                    add_inline(paragraph, item)
            else:
                for item in block.items:
                    paragraph = document.add_paragraph(style="List Bullet")
                    add_inline(paragraph, item)
        elif block.kind == "note":
            render_note(document, block.note_kind, block.text)
        elif block.kind == "roles":
            render_roles(document, block.items)
        elif block.kind == "aag":
            render_box(document, "At a glance", block.items, AAG_BG)
        elif block.kind == "steps":
            render_box(document, "How to", block.items, STEPS_BG)
        elif block.kind == "toc":
            add_toc_field(document)
        elif block.kind == "table":
            render_table(document, block.items, block.rows)
        elif block.kind == "image":
            render_image_block(
                document,
                block.image_path,
                block.image_caption,
                height_pt=block.image_height_pt,
            )
        elif block.kind == "highlight":
            render_highlight_block(
                document,
                block.image_path,
                block.image_caption,
                height_pt=block.image_height_pt,
            )


# --------------------------------- Build ---------------------------------- #


def build() -> Path:
    if not CONTENT_DIR.exists():
        raise SystemExit(f"Content directory missing: {CONTENT_DIR}")
    document = Document()
    configure_page_layout(document)
    configure_styles(document)

    for index, filename in enumerate(SECTIONS):
        md_path = CONTENT_DIR / filename
        if not md_path.exists():
            print(f"[warn] skipping missing content file: {filename}", file=sys.stderr)
            continue
        blocks = parse_markdown(md_path.read_text(encoding="utf-8"))
        render_blocks(document, blocks, section_index=index)

    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    output = build()
    print(f"Wrote: {output}")
