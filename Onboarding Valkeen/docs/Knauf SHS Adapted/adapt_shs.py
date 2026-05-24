"""Apply Knauf customizations to the SHS PM training template.

The script is config-driven; everything is JSON. Three independent steps:

1. **Text replacements** — ``config/text_replacements.json`` swaps
   customer-specific strings (e.g. "Siemens Healthineers" -> "Knauf").

2. **Image replacements** — ``config/image_mapping.json`` overwrites
   selected ``word/media/imageN.png`` entries inside the DOCX zip with
   Knauf screenshots placed under ``screenshots/``. The original
   anchor/size/position in ``document.xml`` is preserved, so the new
   image takes the exact same slot.

3. **Highlight shapes** — ``config/highlights.json`` adds Word-native
   floating rectangles (DrawingML ``wps:wsp``) anchored to paragraphs
   that contain the screenshots. They have a transparent fill and a
   coloured border, so they render as classic "highlight boxes".

   Because they are real Word shapes, they can be **selected, dragged,
   resized, recoloured** directly in Microsoft Word once the document
   is opened.

Usage::

    # Inventory of all media files inside the SHS DOCX
    python adapt_shs.py --list-images

    # Rebuild the Knauf DOCX from scratch (source -> output)
    python adapt_shs.py

The output ``Knauf_PM_Training.docx`` is regenerated on every run, so
the script is safe to re-run after editing any of the three configs.
"""

from __future__ import annotations

import argparse
import io
import json
import shutil
import sys
import uuid
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from PIL import Image
from lxml import etree


HERE = Path(__file__).resolve().parent
SOURCE_DOCX = HERE.parent / "_shs-source" / "SHS_PM_Original.docx"
OUTPUT_DOCX = HERE / "Knauf_PM_Training.docx"
SCREENSHOTS = HERE / "screenshots"
CFG = HERE / "config"


# --------------------------------------------------------------------------- #
# Text replacements                                                           #
# --------------------------------------------------------------------------- #


def _replace_in_paragraphs(paragraphs, replacements: list[dict]) -> int:
    n = 0
    for p in paragraphs:
        for src, dst in [(r["from"], r["to"]) for r in replacements]:
            # 1. Try cheap run-level replace first (preserves formatting).
            for r in p.runs:
                if src in r.text:
                    r.text = r.text.replace(src, dst)
                    n += 1
            # 2. Fall back to paragraph-level when the match is split
            #    across multiple runs.
            full = "".join(r.text for r in p.runs)
            if src in full and not any(src in r.text for r in p.runs):
                new_text = full.replace(src, dst)
                if p.runs:
                    p.runs[0].text = new_text
                    for r in p.runs[1:]:
                        r.text = ""
                    n += 1
    return n


def apply_text_replacements(doc, cfg_path: Path) -> int:
    cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
    rules = cfg.get("replacements", [])
    if not rules:
        return 0
    n = _replace_in_paragraphs(doc.paragraphs, rules)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                n += _replace_in_paragraphs(cell.paragraphs, rules)
    for section in doc.sections:
        n += _replace_in_paragraphs(section.header.paragraphs, rules)
        n += _replace_in_paragraphs(section.footer.paragraphs, rules)
    return n


# --------------------------------------------------------------------------- #
# Image replacements (DOCX = ZIP, media lives under word/media/)              #
# --------------------------------------------------------------------------- #


def _fit_to_original_size(new_path: Path, original_bytes: bytes,
                          pad_color: tuple[int, int, int] = (255, 255, 255)) -> bytes:
    """Resize ``new_path`` so its pixel dimensions match the original
    image (read from ``original_bytes``). Uses an *aspect-fit* strategy
    with padding so the new content never gets distorted.

    Returns the resulting PNG bytes.
    """
    orig = Image.open(io.BytesIO(original_bytes))
    target_w, target_h = orig.size
    new_im = Image.open(new_path).convert("RGB")
    # aspect-fit
    new_im.thumbnail((target_w, target_h), Image.LANCZOS)
    canvas = Image.new("RGB", (target_w, target_h), pad_color)
    off_x = (target_w - new_im.width) // 2
    off_y = (target_h - new_im.height) // 2
    canvas.paste(new_im, (off_x, off_y))
    buf = io.BytesIO()
    canvas.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


def apply_image_replacements(docx_path: Path, cfg_path: Path) -> int:
    cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
    mapping = cfg.get("replacements", {})
    if not mapping:
        return 0

    tmp = docx_path.with_suffix(".tmp.docx")
    n = 0
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                name = Path(item).name
                if item.startswith("word/media/") and name in mapping:
                    new_path = HERE / mapping[name]
                    if not new_path.exists():
                        raise FileNotFoundError(
                            f"image_mapping references missing file: {new_path}"
                        )
                    original_bytes = zin.read(item)
                    fitted = _fit_to_original_size(new_path, original_bytes)
                    zout.writestr(item, fitted)
                    print(f"  [img]  {name:20} <- {mapping[name]}")
                    n += 1
                else:
                    zout.writestr(item, zin.read(item))
    docx_path.unlink()
    tmp.rename(docx_path)
    return n


# --------------------------------------------------------------------------- #
# Highlight shapes (Word-native, editable after document is opened)           #
# --------------------------------------------------------------------------- #


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"


def _cm_to_emu(cm: float) -> int:
    return int(cm * 360000)


def _build_rect_xml(
    *, posH_emu: int, posV_emu: int, cx_emu: int, cy_emu: int,
    color: str, line_w_emu: int, shape_id: int, label: str,
) -> etree._Element:
    """Build a floating DrawingML rectangle (`wps:wsp`) anchored to the
    paragraph it is inserted in."""

    xml = f'''<w:r xmlns:w="{W_NS}" xmlns:wp="{WP_NS}" xmlns:a="{A_NS}"
       xmlns:wps="{WPS_NS}" xmlns:mc="{MC_NS}">
  <mc:AlternateContent>
    <mc:Choice Requires="wps">
      <w:drawing>
        <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
                   simplePos="0" relativeHeight="{251660000 + shape_id}"
                   behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="column">
            <wp:posOffset>{posH_emu}</wp:posOffset>
          </wp:positionH>
          <wp:positionV relativeFrom="paragraph">
            <wp:posOffset>{posV_emu}</wp:posOffset>
          </wp:positionV>
          <wp:extent cx="{cx_emu}" cy="{cy_emu}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:wrapNone/>
          <wp:docPr id="{shape_id}" name="Highlight {label}"/>
          <wp:cNvGraphicFramePr/>
          <a:graphic>
            <a:graphicData uri="{WPS_NS}">
              <wps:wsp>
                <wps:cNvSpPr/>
                <wps:spPr>
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="{cx_emu}" cy="{cy_emu}"/>
                  </a:xfrm>
                  <a:prstGeom prst="roundRect"><a:avLst><a:gd name="adj" fmla="val 8000"/></a:avLst></a:prstGeom>
                  <a:noFill/>
                  <a:ln w="{line_w_emu}">
                    <a:solidFill><a:srgbClr val="{color}"/></a:solidFill>
                  </a:ln>
                </wps:spPr>
                <wps:bodyPr/>
              </wps:wsp>
            </a:graphicData>
          </a:graphic>
        </wp:anchor>
      </w:drawing>
    </mc:Choice>
  </mc:AlternateContent>
</w:r>'''
    return etree.fromstring(xml)


def apply_highlights(doc, cfg_path: Path) -> int:
    cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
    highlights = cfg.get("highlights", [])
    if not highlights:
        return 0

    n = 0
    for idx, h in enumerate(highlights, start=1):
        p_idx = h["paragraph_index"]
        if p_idx >= len(doc.paragraphs):
            print(f"  [warn] paragraph_index {p_idx} out of range; skipping")
            continue
        p = doc.paragraphs[p_idx]
        rect = _build_rect_xml(
            posH_emu=_cm_to_emu(h.get("x_cm", 0.0)),
            posV_emu=_cm_to_emu(h.get("y_cm", 0.0)),
            cx_emu=_cm_to_emu(h["w_cm"]),
            cy_emu=_cm_to_emu(h["h_cm"]),
            color=h.get("color", "F2B807").lstrip("#").upper(),
            line_w_emu=int(h.get("line_width_pt", 2.0) * 12700),  # 1 pt = 12700 EMU
            shape_id=1000 + idx,
            label=h.get("label", str(idx)),
        )
        p._p.append(rect)
        n += 1
    return n


# --------------------------------------------------------------------------- #
# Inventory                                                                   #
# --------------------------------------------------------------------------- #


def list_images() -> None:
    """Print a table of every image embedded in the SHS DOCX so you can
    decide which ones to map to Knauf screenshots."""
    print(f"# Inventory of images inside {SOURCE_DOCX.name}\n")
    print(f"{'idx':>3} {'media':18} {'size_kb':>7}  {'pixels':>12}  notes")
    print("-" * 72)
    with zipfile.ZipFile(SOURCE_DOCX) as z:
        media = sorted(n for n in z.namelist() if n.startswith("word/media/"))
        for i, m in enumerate(media):
            data = z.read(m)
            name = Path(m).name
            kb = len(data) // 1024
            dims = ""
            try:
                im = Image.open(io.BytesIO(data))
                dims = f"{im.size[0]}x{im.size[1]}"
            except Exception:
                pass
            print(f"{i:>3} {name:18} {kb:>5} KB  {dims:>12}")


# --------------------------------------------------------------------------- #
# Main                                                                        #
# --------------------------------------------------------------------------- #


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--list-images", action="store_true", help="Print image inventory and exit")
    args = parser.parse_args()

    if args.list_images:
        list_images()
        return

    # 1. Clone the source template fresh.
    print(f"[step] cloning  {SOURCE_DOCX.name}  ->  {OUTPUT_DOCX.name}")
    shutil.copy(SOURCE_DOCX, OUTPUT_DOCX)

    # 2. Image replacements at the ZIP level.
    img_cfg = CFG / "image_mapping.json"
    if img_cfg.exists():
        n_img = apply_image_replacements(OUTPUT_DOCX, img_cfg)
        print(f"[step] images replaced: {n_img}")

    # 3. Text replacements + highlight shapes (both via python-docx /
    #    lxml, so we must reopen the document after the ZIP edit).
    doc = Document(OUTPUT_DOCX)
    txt_cfg = CFG / "text_replacements.json"
    if txt_cfg.exists():
        n_txt = apply_text_replacements(doc, txt_cfg)
        print(f"[step] text replacements: {n_txt}")
    hl_cfg = CFG / "highlights.json"
    if hl_cfg.exists():
        n_hl = apply_highlights(doc, hl_cfg)
        print(f"[step] highlight shapes added: {n_hl}")
    doc.save(OUTPUT_DOCX)
    print(f"[ok]   {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
